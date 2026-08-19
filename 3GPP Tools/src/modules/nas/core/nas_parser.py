import json
import logging
from pathlib import Path
import re
from typing import Any, Callable, Dict, List, Optional, Tuple
import docx
from docx.table import Table

from modules.specifications.utils.utils import file_version_to_version


class NASDocxParser:
    """Parses Clause 8 message tables and Clause 9 IE definitions from TS 24.501 .docx specifications."""

    def __init__(self, docx_path: Path):
        self.docx_path = Path(docx_path)
        self.logger = logging.getLogger(__name__)

    def extract_version_from_filename(self) -> str:
        """
        Uses modules.specifications.utils.file_version_to_version to parse
        filenames such as '24501-j30.docx' into '19.3.0'.
        """
        stem = self.docx_path.stem  # e.g., '24501-j30'
        match = re.search(r"-([a-zA-Z0-9]{3})$", stem)
        if match:
            parsed_ver = file_version_to_version(match.group(1))
            if parsed_ver:
                return parsed_ver
        return stem

    def parse(
        self, progress_callback: Optional[Callable[[str, int], None]] = None
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Scans the document for Clause 8 tables and Clause 9 IE headings."""
        if not self.docx_path.exists():
            raise FileNotFoundError(f"Specification file not found: {self.docx_path}")

        if progress_callback:
            progress_callback("Opening document structure...", 5)
        doc = docx.Document(str(self.docx_path))

        messages: List[Dict[str, Any]] = []
        ie_definitions: List[Dict[str, Any]] = []

        total_tables = len(doc.tables)
        if progress_callback:
            progress_callback(
                f"Found {total_tables} tables. Scanning Clause 8...", 15
            )

        # Regex patterns for Clause 8 table captions and Clause 9 subclauses
        caption_pattern = re.compile(
            r"Table\s+(8\.\d+(?:\.\d+)+)\s*:\s*(.+?)(?:\s+message\s+content)?$",
            re.IGNORECASE,
        )
        ie_heading_pattern = re.compile(r"^(9\.11(?:\.\d+)+)\s+(.+)$")

        # 1. Parse Clause 8 Message Content Tables
        for idx, table in enumerate(doc.tables):
            caption = self._find_table_caption(table)
            if caption:
                match = caption_pattern.search(caption)
                if match:
                    clause = match.group(1).strip()
                    name = match.group(2).strip()
                    name = re.sub(r"(?i)\s+message\s+content", "", name).strip()

                    ies = self._parse_clause_8_table(table)
                    if ies:
                        messages.append({
                            "clause": clause,
                            "message_name": name,
                            "table_caption": caption,
                            "ies": ies,
                        })

            if idx % 25 == 0 and progress_callback:
                progress = 15 + int((idx / max(1, total_tables)) * 65)
                progress_callback(f"Parsed {idx}/{total_tables} tables...", progress)

        # 2. Parse Clause 9 Information Element Definitions
        if progress_callback:
            progress_callback("Extracting Clause 9 definitions...", 85)

        for p in doc.paragraphs:
            text = p.text.strip()
            match = ie_heading_pattern.match(text)
            if match:
                cl = match.group(1).strip()
                ie_name = match.group(2).strip()
                ie_definitions.append({
                    "clause": cl,
                    "ie_name": ie_name,
                    "raw_description": f"Definition and coding for '{ie_name}' (Clause {cl}).",
                    "structure_table": json.dumps([]),
                })

        if progress_callback:
            progress_callback(
                f"Extracted {len(messages)} messages and {len(ie_definitions)} IE definitions.",
                100,
            )

        return messages, ie_definitions

    def _find_table_caption(self, table: Table) -> str:
        """Finds preceding caption paragraphs matching 'Table 8.'."""
        prev_p = table._element.getprevious()
        while prev_p is not None:
            if prev_p.tag.endswith("p"):
                text = "".join(prev_p.itertext()).strip()
                if text.startswith("Table 8."):
                    return text
            prev_p = prev_p.getprevious()
        return ""

    def _parse_clause_8_table(self, table: Table) -> List[Dict[str, str]]:
        """Parses the 6 standard columns of a Clause 8 message table."""
        if not table.rows:
            return []

        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        if not any("information element" in h for h in header_cells):
            return []

        ies = []
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 6:
                continue

            iei = cells[0]
            ie_name = cells[1]
            type_ref = cells[2]
            presence = cells[3]
            fmt = cells[4]
            length = cells[5]

            if not ie_name or "information element" in ie_name.lower():
                continue

            ies.append({
                "iei": iei,
                "information_element": ie_name,
                "type_reference": type_ref,
                "presence": presence,
                "format": fmt,
                "length": length,
            })
        return ies